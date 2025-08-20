import boto3
import os
import io
import json
import docx
import pandas as pd
import pdfplumber

REGION =  os.environ.get('AWS_REGION')

s3 = boto3.client('s3', region_name = REGION)
sqs = boto3.client('sqs', region_name = REGION)
QUEUE_URL = os.environ['SQS_QUEUE_URL']

def lambda_handler(event, context):
    # Get bucket and key from event (e.g. S3 trigger or API)
    bucket = event['bucket']
    key = event['key']
    file_type = key.split('.')[-1].lower()

    file_obj = s3.get_object(Bucket=bucket, Key=key)
    file_bytes = file_obj['Body'].read()
    content = extract_text(file_bytes, file_type)
    chunks = smart_chunk(content, file_type)

    for i, chunk in enumerate(chunks):
        sqs.send_message(
            QueueUrl=QUEUE_URL,
            MessageBody=json.dumps({
                "chunk_index": i,
                "chunk": chunk,
                "source_file": key
            })
        )

    return {"chunks_sent": len(chunks), "file": key}


def extract_text(file_bytes, file_type):
    if file_type == 'pdf':
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            texts, tables = [], []
            for page in pdf.pages:
                texts.append(page.extract_text() or "")
                table = page.extract_table()
                if table:
                    table_text = "\n".join([", ".join(row) for row in table])
                    tables.append("TABLE:\n" + table_text)
            return "\n\n".join(texts + tables)

    elif file_type == 'docx':
        doc = docx.Document(io.BytesIO(file_bytes))
        texts = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
        for table in doc.tables:
            rows = [" , ".join([cell.text.strip() for cell in row.cells]) for row in table.rows]
            texts.append("TABLE:\n" + "\n".join(rows))
        return "\n\n".join(texts)

    elif file_type == 'txt':
        return file_bytes.decode('utf-8')

    elif file_type in ['csv', 'xlsx']:
        if file_type == 'csv':
            df = pd.read_csv(io.BytesIO(file_bytes))
        else:
            df = pd.read_excel(io.BytesIO(file_bytes))
        return df.to_csv(index=False)

    else:
        return "Unsupported file type"

def smart_chunk(text, file_type, max_words=200):
    if file_type in ['csv', 'xlsx']:
        rows = text.strip().split('\n')
        return ["\n".join(rows[i:i+10]) for i in range(0, len(rows), 10)]

    paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
    chunks, current = [], []

    for para in paragraphs:
        current.append(para)
        if len(" ".join(current).split()) > max_words:
            chunks.append("\n\n".join(current))
            current = []

    if current:
        chunks.append("\n\n".join(current))

    return chunks

